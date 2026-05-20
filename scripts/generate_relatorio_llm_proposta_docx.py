from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_kv_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"

    for k, v in items:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v


def add_component_table(doc: Document) -> None:
    headers = ["Componente", "Responsabilidade", "Observações"]
    rows = [
        [
            "Supabase Edge Function `ai-orchestrator`",
            "Orquestra conversação, tool-calling, políticas (RBAC/PII), memória e RAG",
            "Mantém streaming SSE; apenas troca o endpoint do provedor LLM.",
        ],
        [
            "LLM self-host (vLLM)",
            "Geração de linguagem + tool-calling (OpenAI-compatible)",
            "Modelo: Qwen2.5-32B-Instruct; endpoint /v1/chat/completions.",
        ],
        [
            "RAG (pgvector no Supabase)",
            "Embeddings + busca vetorial em conteúdo interno",
            "Melhora consistência e reduz alucinação; permite citações/links internos.",
        ],
        [
            "Load Balancer + Cloud Armor",
            "Exposição pública com proteção (WAF, rate limit, regras)",
            "Entrada HTTPS; health checks; bloqueio de abuso.",
        ],
        [
            "Observabilidade",
            "Uptime, latência (p50/p95), erros, tokens, saturação GPU",
            "Alertas para SLO 24/7 e detecção de degradação.",
        ],
        [
            "Fallback SaaS (opcional, recomendado)",
            "Continuidade do chat quando self-host estiver indisponível",
            "Usar OpenAI/Gemini somente em incidentes (circuit breaker).",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def add_risks_table(doc: Document) -> None:
    headers = ["Risco", "Impacto", "Mitigação"]
    rows = [
        [
            "Picos de concorrência desconhecidos",
            "Fila/latência alta (p95) e timeouts",
            "Começar com 2 VMs + rate limit; monitorar concorrência; ajustar batching/limites e escalar horizontalmente.",
        ],
        [
            "Exposição pública sem proteção adequada",
            "Abuso/custos, vazamento de chave, negação de serviço",
            "Cloud Armor + rate limit + WAF; API key/HMAC; limites de tokens/payload; logs e alertas.",
        ],
        [
            "Degradação do modelo em conversa geral",
            "Experiência ruim / respostas incoerentes",
            "RAG + tool-first; testes de qualidade; prompt tuning; fallback SaaS em incidentes.",
        ],
        [
            "Custos de GPU e ociosidade",
            "Custo fixo elevado em baixa demanda",
            "Ajustar horários e capacidade; autoscaling futuro (migração para GKE) se necessário.",
        ],
        [
            "PII/LGPD no histórico",
            "Risco legal e de privacidade",
            "Redação de PII antes de persistir memória longa; políticas de retenção; auditoria de acesso.",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_dir = repo_root / "docs" / "entregaveis"
    out_dir.mkdir(parents=True, exist_ok=True)

    today = date.today().isoformat()
    out_path = out_dir / f"Relatorio_Proposta_Assistente_IA_{today}.docx"

    doc = Document()

    add_title(doc, "Câmara na Mão — Relatório: Proposta de Assistente IA (Self-hosted)")
    add_subtitle(doc, f"Data: {today}")
    doc.add_paragraph()

    add_heading(doc, "1. Objetivo")
    doc.add_paragraph(
        "Definir a proposta técnica para operar o assistente com provedor LLM self-hosted, "
        "mantendo conversação geral, contextualização, chatbot de serviços e RAG, com operação 24/7."
    )

    add_heading(doc, "2. Decisões tomadas")
    add_kv_table(
        doc,
        [
            ("Região (GCP)", "southamerica-east1"),
            ("Hospedagem", "Compute Engine (GPU)"),
            ("SLA", "24/7"),
            ("Qualidade", "Prioridade máxima (32B)"),
            ("Modelo alvo", "Qwen2.5-32B-Instruct"),
            ("Serving", "vLLM (API compatível com OpenAI + streaming)"),
            ("Exposição", "Público com proteção (HTTPS LB + Cloud Armor + Auth na API)"),
            ("Alta disponibilidade", "2 VMs (zonas diferentes) atrás de Load Balancer"),
            ("Fallback SaaS", "Opcional e recomendado (somente contingência via circuit breaker)"),
        ],
    )

    add_heading(doc, "3. Compute Engine vs GKE (comparativo resumido)")
    add_bullets(
        doc,
        [
            "Compute Engine: melhor para começar rápido e com menos operação (recomendado para o 1º mês).",
            "GKE: melhor para escala e rollouts automáticos, porém com mais complexidade/custo operacional.",
            "Estratégia sugerida: iniciar em Compute Engine com API OpenAI-compatible; migrar para GKE se/quando houver necessidade de autoscaling.",
        ],
    )

    add_heading(doc, "4. Arquitetura proposta (alto nível)")
    add_code_block(
        doc,
        [
            "Web/Mobile",
            "  -> Supabase Edge Function (ai-orchestrator)",
            "     -> LLM self-host (vLLM, OpenAI-compatible, streaming)",
            "     -> Tools (agenda, noticias, vereadores, services, etc.)",
            "     -> RAG (pgvector no Supabase) + Memória (curto/longa)",
            "",
            "Internet -> HTTPS Load Balancer -> Cloud Armor -> (VM GPU A / VM GPU B) -> vLLM",
        ],
    )

    add_heading(doc, "5. Componentes e responsabilidades")
    add_component_table(doc)

    add_heading(doc, "6. Segurança (público com proteção)")
    add_bullets(
        doc,
        [
            "Load Balancer com HTTPS + certificado gerenciado.",
            "Cloud Armor (WAF managed rules/OWASP) + rate limiting.",
            "Autenticação obrigatória na API do LLM: API Key (Bearer) e/ou HMAC com timestamp (anti-replay).",
            "Limites de payload e tokens por requisição; timeouts curtos; logs e alertas.",
        ],
    )

    add_heading(doc, "7. Operação 24/7 (alta disponibilidade)")
    add_bullets(
        doc,
        [
            "2 VMs GPU em zonas diferentes na região southamerica-east1 (reduz risco de falha zonal).",
            "Health checks no LB e remoção automática de instância degradada.",
            "Atualização rolling: atualizar 1 VM por vez para evitar downtime.",
            "Observabilidade: p50/p95, erros, saturação GPU, tokens/req, filas/timeouts.",
        ],
    )

    add_heading(doc, "8. Fallback SaaS (o que é e por que usar)")
    doc.add_paragraph(
        "Fallback SaaS é um plano de contingência: se o LLM self-host estiver indisponível ou instável, "
        "o orquestrador redireciona temporariamente as requisições para um provedor SaaS (ex.: OpenAI ou Google Gemini) "
        "para manter a disponibilidade 24/7. O fallback deve ser controlado por circuit breaker, limites de custo e "
        "regras de privacidade (ex.: não enviar PII)."
    )
    add_bullets(
        doc,
        [
            "Ativa somente em incidente (timeouts/erros repetidos).",
            "Duração curta (ex.: 5–15 min) e com logs/auditoria.",
            "Desligável por feature flag (ambientes/fluxos sensíveis).",
        ],
    )

    add_heading(doc, "9. Mudanças necessárias no projeto")
    add_bullets(
        doc,
        [
            "Atualizar `supabase/functions/ai-orchestrator/index.ts` para usar `CAMARA_LLM_BASE_URL` / `AI_CHAT_BASE_URL`.",
            "Adicionar secrets/envs: `CAMARA_LLM_BASE_URL` e `CAMARA_LLM_API_KEY` (ou `CAMARA_LLM_HMAC_SECRET`).",
            "Adicionar timeouts, retry curto e circuit breaker no `ai-orchestrator`.",
            "Manter o padrão tool-first + RAG + memória por conversa.",
        ],
    )

    add_heading(doc, "10. Checklist de implantação (go-live)")
    add_numbered(
        doc,
        [
            "Criar 2 VMs GPU em zonas diferentes (southamerica-east1) e instalar vLLM + modelo.",
            "Subir serviço vLLM como systemd (auto-restart) e expor endpoint /v1/chat/completions.",
            "Criar HTTPS Load Balancer com backend para as 2 VMs e health check.",
            "Configurar Cloud Armor (WAF + rate limiting) e regras adicionais (geo/allowlist, se necessário).",
            "Definir autenticação (API Key e/ou HMAC) e aplicar no proxy/API.",
            "Configurar secrets no Supabase e alterar o `ai-orchestrator` para o novo endpoint.",
            "Habilitar logs/alertas (uptime, p95, erros, consumo).",
            "Rodar teste de carga (picos) e ajustar limites (tokens, concurrency, timeouts).",
            "Planejar fallback SaaS (opcional) com circuit breaker e limites de custo.",
        ],
    )

    add_heading(doc, "11. Riscos e mitigações")
    add_risks_table(doc)

    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()

