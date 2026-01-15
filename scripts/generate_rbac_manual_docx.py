from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, lines: list[str]) -> None:
    # Simple "code block" approximation: monospaced font + smaller size.
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_matrix_table(doc: Document) -> None:
    headers = ["Funcionalidade", "Cidadão", "Cidadão Engajado", "Gestor", "Admin"]
    rows = [
        ["Criar manifestações", "✅", "✅", "✅", "✅"],
        ["Ver próprias manifestações", "✅", "✅", "✅", "✅"],
        ["Avaliar serviços", "✅", "✅", "✅", "✅"],
        ["Buscar serviços próximos", "✅", "✅", "✅", "✅"],
        ["Inscrever-se em audiências", "✅", "✅", "✅", "✅"],
        ["Encaminhar para vereador", "❌", "✅", "✅", "✅"],
        ["Ver dashboards (público)", "❌", "✅", "✅", "✅"],
        ["Criar dashboards", "❌", "✅", "✅", "✅"],
        ["Responder manifestações", "❌", "❌", "✅", "✅"],
        ["Gerenciar triagem", "❌", "❌", "✅", "✅"],
        ["Exportar dados", "❌", "❌", "✅", "✅"],
        ["Configurar sistema", "❌", "❌", "❌", "✅"],
        ["Gerenciar usuários", "❌", "❌", "❌", "✅"],
        ["Acessar logs de auditoria", "❌", "❌", "❌", "✅"],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = value


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_dir = repo_root / "docs" / "entregaveis"
    out_dir.mkdir(parents=True, exist_ok=True)

    today = date.today().isoformat()
    out_path = out_dir / f"Manual_Testes_RBAC_{today}.docx"

    doc = Document()

    add_title(doc, "Câmara na Mão — Manual de Testes (RBAC)")
    add_subtitle(doc, f"Data: {today}")
    doc.add_paragraph()

    add_heading(doc, "1. Objetivo")
    doc.add_paragraph(
        "Este documento descreve como validar as permissões (RBAC) para os perfis "
        "Cidadão, Cidadão Engajado, Gestor e Admin, cobrindo os fluxos do frontend "
        "e as proteções no banco (RLS no Supabase)."
    )

    add_heading(doc, "2. Matriz de permissões (alto nível)")
    add_matrix_table(doc)

    add_heading(doc, "3. Pré-requisitos")
    add_bullets(
        doc,
        [
            "Acesso ao projeto (branch que contém RBAC: pov-dev-mobile ou equivalente).",
            "Docker instalado e rodando (para Supabase local).",
            "Node instalado (para rodar o frontend).",
        ],
    )

    add_heading(doc, "4. Subindo o Supabase local")
    doc.add_paragraph("No diretório raiz do projeto, rode:")
    add_code_block(
        doc,
        [
            r"cd C:\Projetos\camana-na-mao",
            r"npx supabase@latest start --exclude studio",
            r"npx supabase@latest status",
        ],
    )

    add_heading(doc, "5. Configurando o frontend web para o Supabase local")
    doc.add_paragraph("Crie/ajuste o arquivo .env.local (não commitar) com:")
    add_code_block(
        doc,
        [
            "VITE_SUPABASE_URL=http://127.0.0.1:54321",
            "VITE_SUPABASE_PUBLISHABLE_KEY=<SUA_SB_PUBLISHABLE_KEY>",
        ],
    )
    doc.add_paragraph("Suba o frontend:")
    add_code_block(doc, [r"npm run dev"])

    add_heading(doc, "6. Criando usuários de teste")
    add_numbered(
        doc,
        [
            "Crie 4 contas no frontend (emails diferentes), uma por perfil: U1=Cidadão, U2=Cidadão Engajado, U3=Gestor, U4=Admin.",
            "Faça login/logout após promoções para recarregar permissões no frontend.",
        ],
    )

    add_heading(doc, "7. Descobrindo o UUID e promovendo roles")
    doc.add_paragraph("Liste os últimos usuários (PowerShell):")
    add_code_block(
        doc,
        [
            r"docker exec supabase_db_vzkwkcypkfrpfhhsghwn psql -U postgres -d postgres -c ""select id, email, created_at from auth.users order by created_at desc limit 10;""",
        ],
    )
    doc.add_paragraph("Para promover um role (repita para cada usuário/role):")
    add_code_block(
        doc,
        [
            r"docker exec supabase_db_vzkwkcypkfrpfhhsghwn psql -U postgres -d postgres -c ""insert into public.user_roles (user_id, role) values ('<USER_UUID>', '<ROLE>'::public.app_role) on conflict do nothing;""",
        ],
    )
    add_bullets(
        doc,
        [
            "U1 (Cidadão): não adicionar nada (deve ficar apenas cidadao).",
            "U2 (Cidadão Engajado): adicionar role cidadao_engajado.",
            "U3 (Gestor): adicionar role gestor.",
            "U4 (Admin): adicionar role admin.",
        ],
    )

    add_heading(doc, "8. Roteiro de testes por funcionalidade")
    doc.add_paragraph(
        "Abaixo, valide sempre em 2 camadas: (1) a UI deve bloquear/permitir e (2) o banco deve bloquear/permitir via RLS."
    )

    add_heading(doc, "8.1 Manifestações (criar / ver próprias)", level=3)
    add_bullets(
        doc,
        [
            "Criar manifestações (Urban/Transport): esperado ✅ para todos.",
            "Ver próprias manifestações: esperado ✅ para todos.",
        ],
    )

    add_heading(doc, "8.2 Encaminhar para vereador", level=3)
    add_bullets(
        doc,
        [
            "U1 (Cidadão): esperado ❌ (UI bloqueia + RLS bloqueia INSERT em council_member_referrals).",
            "U2/U3/U4: esperado ✅ (consegue concluir o fluxo).",
        ],
    )
    doc.add_paragraph("Validação opcional no banco:")
    add_code_block(
        doc,
        [
            r"docker exec supabase_db_vzkwkcypkfrpfhhsghwn psql -U postgres -d postgres -c ""select user_id, council_member_name, status, created_at from public.council_member_referrals order by created_at desc limit 10;""",
        ],
    )

    add_heading(doc, "8.3 Dashboards (ver público / criar)", level=3)
    add_bullets(
        doc,
        [
            "U1 (Cidadão): esperado ❌ (sem acesso a /paineis e /paineis/criar).",
            "U2/U3/U4: esperado ✅ (acessa /paineis e consegue criar).",
        ],
    )

    add_heading(doc, "8.4 Admin: responder manifestações e triagem", level=3)
    add_bullets(
        doc,
        [
            "Responder manifestações: esperado ✅ apenas para Gestor/Admin.",
            "Gerenciar triagem (ver tudo / atualizar status): esperado ✅ apenas para Gestor/Admin.",
        ],
    )

    add_heading(doc, "8.5 Exportação de dados", level=3)
    add_bullets(
        doc,
        [
            "U3/U4: esperado ✅ (criar/visualizar export_logs).",
            "U1/U2: esperado ❌.",
        ],
    )

    add_heading(doc, "8.6 Configurar sistema / Gerenciar usuários / Auditoria", level=3)
    add_bullets(
        doc,
        [
            "Apenas Admin: esperado ✅.",
            "Gestor: esperado ❌ (rotas admin-only).",
            "Cidadão/Engajado: esperado ❌.",
        ],
    )

    add_heading(doc, "9. Evidências recomendadas")
    add_bullets(
        doc,
        [
            "Print do Cidadão bloqueado em /paineis e no fluxo de encaminhamento.",
            "Print do Cidadão Engajado acessando /paineis e concluindo encaminhamento.",
            "Print do Gestor acessando triagem/respostas (admin).",
            "Print do Admin acessando gestão de usuários e logs de auditoria.",
        ],
    )

    add_heading(doc, "10. Troubleshooting (quando 'roles não aparecem')")
    add_bullets(
        doc,
        [
            "Faça logout/login (ou hard refresh) para recarregar permissões no frontend.",
            "Confirme que o frontend está apontando para o mesmo Supabase (URL e key).",
            "No Render: faça 'Clear build cache & deploy' se houver suspeita de bundle antigo.",
        ],
    )

    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()

