from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def add_heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_dir = repo_root / "docs" / "entregaveis"
    out_dir.mkdir(parents=True, exist_ok=True)

    today = date.today().isoformat()
    out_path = out_dir / f"Entregaveis_Camara_na_Mao_{today}.docx"

    doc = Document()

    add_title(doc, "Câmara na Mão — Entregáveis (PoV Mobile + Deploy)")
    add_subtitle(doc, f"Data: {today}")
    doc.add_paragraph()

    add_heading(doc, "Resumo executivo")
    doc.add_paragraph(
        "Consolidamos um PoV mobile (Expo/React Native) para testes em Android, com capacidade de abrir o frontend web "
        "publicado e realizar diagnósticos de conectividade/serviços. Também formalizamos padrões e documentação "
        "arquitetural (ADRs, variáveis de ambiente, lint/format e aliases)."
    )

    add_heading(doc, "Entregas realizadas")
    add_bullets(
        doc,
        [
            "Mobile (React Native/Expo) criado no diretório mobile/ para testes rápidos no Android (Expo Go e APK via EAS).",
            "Aba API no mobile para diagnóstico: campo de URL + botão “Testar API”, com timeout e mensagens claras.",
            "Supabase local via Docker + Supabase CLI (supabase start --exclude studio) para desenvolvimento/testes locais.",
            "Aba Frontend no mobile com WebView para renderizar o frontend web por URL (ambiente público no Render).",
            "Frontend publicado no Render como Static Site, com URL pública para testes externos.",
            "ADRs adicionados (docs/adr/) documentando decisões arquiteturais e trade-offs.",
            "Documentação de variáveis de ambiente (docs/ENVIRONMENT_VARIABLES.md) + env.example; .env ignorado no Git.",
            "ESLint + Prettier configurados (scripts e docs) com integração sem conflito (eslint-config-prettier).",
            "Aliases de importação padronizados: web (@/* → src/*) e mobile (@/* → mobile/src/*).",
            "Branch pov-dev-mobile atualizada com ux-ui (merge) e sincronizada no repositório original e no pessoal.",
        ],
    )

    add_heading(doc, "Links úteis")
    add_bullets(
        doc,
        [
            "Ambiente web (Render): https://camara-na-mao.onrender.com/welcome",
            "Repositório (branch pov-dev-mobile): https://github.com/MTechnologia/camana-na-mao/tree/pov-dev-mobile",
        ],
    )

    add_heading(doc, "Observações")
    add_bullets(
        doc,
        [
            "Para uso por terceiros fora da rede local, o mobile deve apontar para URL pública (HTTPS).",
            "O APK de Android deve ser reinstalado ao atualizar ícone/splash ou URL padrão.",
            "Chaves/segredos não devem ser commitados; usar variáveis no Render e arquivos locais .env (ignorados).",
        ],
    )

    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()

