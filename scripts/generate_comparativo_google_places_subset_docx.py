"""Gera docs/comparativo-google-places-subset-cohort.docx (subconjunto Marco C + Google Places)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "comparativo-google-places-subset-cohort.docx"

# Resultado da 2.ª query em diagnostic-google-places-subset-cohort.sql (Marco C; inclui hospital)
N_TOTAL = 80_079
USD_APPROX = "1.976,98"
BRL_APPROX = "10.873,36"


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h0 = doc.add_heading(
        "Google Places — subconjunto por coorte (Marco C)",
        level=0,
    )
    h0.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph(
        "Estimativa de custo (Place Details) para o recorte definido no comparativo geral "
        "de equipamentos (secção 5), com base na query "
        "scripts/sql/diagnostic-google-places-subset-cohort.sql."
    )

    p = doc.add_paragraph()
    p.add_run("Câmbio ilustrativo: ").bold = True
    p.add_run("R$ 5,50 por US$ 1,00.")

    p2 = doc.add_paragraph()
    p2.add_run("Google Places (estimativa): ").bold = True
    p2.add_run(
        "1 chamada Place Details (Enterprise + Atmosphere) por POI; "
        "US$ 25 / 1.000 após 1.000 grátis/mês; sem ponto de ônibus; sem Popular times."
    )

    doc.add_heading("1. O que entra no subconjunto", level=1)
    _add_table(
        doc,
        ["Área", "Critério (resumo)"],
        [
            [
                "Metrô",
                "transit_station + estacao_metro / estacao_metro_projeto",
            ],
            [
                "Trem",
                "transit_station + estacao_trem / estacao_trem_projeto",
            ],
            [
                "Terminal de ônibus",
                "transit_station + terminal_onibus",
            ],
            [
                "Escola (4 camadas GeoSampa)",
                "school + ensino_fundamental_medio, educacao_infantil, ensino_tecnico, senai_sesi_senac (ver §1.1)",
            ],
            [
                "Saúde (hospitais, UPAs e afins no mapa)",
                "service_type = hospital + source_layer GeoSampa (ver §1.2)",
            ],
            [
                "Parque, centro esportivo, CEU, teatro/cinema, biblioteca, museu",
                "service_type: park, sports_center, ceu, theater, library, museum",
            ],
        ],
    )

    doc.add_heading("1.1. Escolas — quatro camadas GeoSampa", level=2)
    doc.add_paragraph(
        "Entram apenas POIs com service_type = school e source_layer igual a um dos valores da tabela "
        "(WFS GeoSampa → public_services). Fora do recorte: rede_privada, educacao_outros, escolas só da API "
        "Escola Aberta sem layer compatível, etc."
    )
    _add_table(
        doc,
        ["source_layer", "O que representa (resumo)"],
        [
            [
                "ensino_fundamental_medio",
                "Ensino fundamental e médio rede pública (equipamento_educacao_rede_publica).",
            ],
            [
                "educacao_infantil",
                "Educação infantil rede pública (equipamento_educacao_infantil_rede_publica).",
            ],
            [
                "ensino_tecnico",
                "Ensino técnico rede pública (equipamento_educacao_ensino_tecnico_rede_publica).",
            ],
            [
                "senai_sesi_senac",
                "SENAI / SESI / SENAC (equipamento_educacao_senai_sesi_senac).",
            ],
        ],
    )

    doc.add_heading("1.2. Saúde — service_type hospital (hospitais, UPAs e afins)", level=2)
    doc.add_paragraph(
        "No mapa, hospitais, unidades de urgência/emergência (perfil tipo UPA conforme cadastro GeoSampa) "
        "e unidades afins ficam em service_type = hospital, distinguindo-se pelo source_layer."
    )
    _add_table(
        doc,
        ["source_layer", "O que representa (resumo)"],
        [
            [
                "urgencia_emergencia",
                "Urgência e emergência (inclui perfil tipo UPA / atenção 24h conforme GeoSampa).",
            ],
            ["hospital", "Hospitais."],
            [
                "equipamento_saude_ambulatorios_especializados",
                "Ambulatórios especializados.",
            ],
            [
                "equipamento_saude_saude_mental",
                "Equipamentos de saúde mental.",
            ],
            ["equipamento_ccz", "Centro de Controle de Zoonoses (CCZ)."],
            [
                "equipamento_saude_outros",
                "Outros equipamentos de saúde nesta família.",
            ],
            [
                "equipamento_saude_unidades_dst_aids",
                "Unidades DST / AIDS.",
            ],
        ],
    )
    doc.add_paragraph(
        "Outros source_layer ainda com service_type = hospital aparecem na 1.ª query do SQL como "
        "cohort hospital_source_layer_outros (auditoria)."
    )

    doc.add_heading(
        "2. Contagens por grupo (Marco C, alinhado ao comparativo geral)",
        level=1,
    )
    _add_table(
        doc,
        ["Linha", "N"],
        [
            ["Escola (4 camadas, §1.1)", "57.962"],
            ["Transporte (metrô + trem + terminal)", "13.199"],
            ["Parque", "6.125"],
            ["Centro esportivo", "620"],
            ["CEU", "513"],
            ["Teatro / cinema", "410"],
            ["Biblioteca", "384"],
            ["Museu", "167"],
            ["Saúde (hospital / UPA / afins, §1.2)", "699"],
            ["Total", "80.079"],
        ],
    )
    doc.add_paragraph(
        "Detalhe por coorte (transit_*, school_*, hospital_*, etc.): 1.ª SELECT em "
        "diagnostic-google-places-subset-cohort.sql."
    )

    doc.add_heading("3. Total do subconjunto e custo ilustrativo (Marco C)", level=1)
    p3 = doc.add_paragraph()
    p3.add_run("Fórmula: ").bold = True
    p3.add_run("max(0, N − 1.000) × 0,025 USD; BRL = USD × 5,50 (2.ª query SQL).")

    n_fmt = f"{N_TOTAL:,}".replace(",", ".")
    _add_table(
        doc,
        ["Métrica", "Valor"],
        [
            ["POIs (N)", n_fmt],
            ["USD (≈)", USD_APPROX],
            ["BRL (≈)", BRL_APPROX],
        ],
    )
    doc.add_paragraph(
        f"({n_fmt} − 1.000) × 0,025 = {USD_APPROX} USD; × 5,50 = {BRL_APPROX} BRL."
    )

    foot = doc.add_paragraph(
        "Valores ilustrativos; não substituem fatura Google. "
        "Ver também: comparativo-equipamentos-antes-depois-google-places.md."
    )
    foot.italic = True

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            continue
        if para.alignment is None:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Escrito: {OUT}")


if __name__ == "__main__":
    main()
