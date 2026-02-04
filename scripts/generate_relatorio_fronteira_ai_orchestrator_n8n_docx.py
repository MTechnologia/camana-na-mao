from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_kv_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"

    for k, v in items:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v


def add_boundary_table(doc: Document) -> None:
    headers = ["Tipo", "Responsável", "Quando usar", "Exemplos no projeto"]
    rows = [
        [
            "Síncrono (tempo real)",
            "ai-orchestrator (Edge Function)",
            "Responde ao usuário no chat com baixa latência e streaming (SSE); controla contexto, tools e políticas",
            "supabase/functions/ai-orchestrator; tools de consulta (agenda/notícias/vereadores), RAG, criação transacional (relatos/inscrições)",
        ],
        [
            "Assíncrono (background)",
            "n8n (workflows)",
            "Processos pesados, integrações externas, retries longos, rotinas e enriquecimento pós-evento",
            "supabase/functions/notify-n8n, n8n-webhook, n8n-callback; docs/N8N_INTEGRATION_GUIDE.md",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def add_examples_table(doc: Document) -> None:
    headers = ["Capacidade", "ai-orchestrator (síncrono)", "n8n (assíncrono)"]
    rows = [
        [
            "Conversação geral + streaming",
            "Sim (SSE token a token, UX do chat)",
            "Não recomendado (overhead, difícil streaming consistente)",
        ],
        [
            "RAG (contexto por base interna)",
            "Sim (retrieval + prompt + resposta imediata)",
            "Opcional (pré-processar e curar base; jobs de reindex)",
        ],
        [
            "Tool-calling com validação forte",
            "Sim (schemas, RBAC, retries curtos, circuit breaker)",
            "Parcial (possível, mas menos robusto para tempo real)",
        ],
        [
            "Integrações externas (CRM/ticket/etc.)",
            "Evitar no caminho crítico; só disparo de evento",
            "Sim (ideal: retries longos, idempotência, múltiplas etapas)",
        ],
        [
            "Enriquecimento de relatos (classificar, tags, resumo)",
            "No máximo o essencial para confirmar ao usuário",
            "Sim (melhor: processamento pesado e posterior callback)",
        ],
        [
            "Notificações e rotinas",
            "Apenas criação do registro/trigger",
            "Sim (envio, agendamento, escalonamento, auditoria externa)",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_dir = repo_root / "docs" / "entregaveis"
    out_dir.mkdir(parents=True, exist_ok=True)

    today = date.today().isoformat()
    out_path = out_dir / f"Relatorio_Fronteira_ai-orchestrator_n8n_{today}.docx"

    doc = Document()

    add_title(doc, "Câmara na Mão — Relatório: Fronteira entre ai-orchestrator e n8n")
    add_subtitle(doc, f"Data: {today}")
    doc.add_paragraph()

    add_heading(doc, "1. Objetivo")
    doc.add_paragraph(
        "Definir claramente a separação de responsabilidades entre o **ai-orchestrator** "
        "(orquestrador do chat em tempo real) e o **n8n** (workflows assíncronos), "
        "evitando acoplamento, reduzindo latência e aumentando confiabilidade."
    )

    add_heading(doc, "2. Definições (para quem está lendo pela primeira vez)")
    add_kv_table(
        doc,
        [
            (
                "ai-orchestrator (o que é)",
                "Uma **Supabase Edge Function** que recebe as mensagens do chat e coordena: "
                "contexto/memória, RAG, chamada do LLM, tool-calling e regras (RBAC/PII), "
                "devolvendo resposta em **streaming (SSE)**.",
            ),
            (
                "n8n (o que é)",
                "Uma plataforma de **automação/workflows** para processos assíncronos: "
                "integrações externas, rotinas e etapas longas com retries, sem travar o chat.",
            ),
        ],
    )

    add_heading(doc, "3. Regra de ouro (separação)")
    add_bullets(
        doc,
        [
            "Tudo que precisa responder o usuário **na hora** (UX do chat) → **ai-orchestrator (síncrono)**.",
            "Tudo que pode rodar em **background** (segundos/minutos), com integrações e retries longos → **n8n (assíncrono)**.",
        ],
    )

    add_heading(doc, "4. Tabela de fronteira (resumo)")
    add_boundary_table(doc)

    add_heading(doc, "5. O que deve ficar no ai-orchestrator (síncrono)")
    add_bullets(
        doc,
        [
            "Conversação geral + streaming SSE (resposta token a token).",
            "Roteamento de intenção/jornada (urban_report, transport_report, audiencias, services, etc.).",
            "RAG (buscar contexto relevante e responder com evidência).",
            "Tools de consulta rápida (agenda, notícias, vereadores, serviços próximos, places).",
            "Ações transacionais simples e rápidas (criar relato, registrar inscrição, registrar avaliação), quando necessário para confirmar ao usuário.",
            "RBAC/PII/guardrails e validação de argumentos antes de executar qualquer tool.",
        ],
    )

    add_heading(doc, "6. O que deve ir para n8n (assíncrono)")
    add_bullets(
        doc,
        [
            "Enriquecimento pesado de relatos (classificação, tags, prioridade, deduplicação, sumarização).",
            "Integrações externas (ticketing/CRM/portais), com múltiplas etapas e retries longos.",
            "Notificações e comunicações (email/whatsapp/push), agendamentos e escalonamentos.",
            "Rotinas (ETL, sincronizações, relatórios periódicos, reindex de embeddings).",
        ],
    )

    add_heading(doc, "7. Exemplos práticos (matriz de decisão)")
    add_examples_table(doc)

    add_heading(doc, "8. Fluxos recomendados (padrão operacional)")
    add_numbered(
        doc,
        [
            "Chat → ai-orchestrator coleta dados → cria registro essencial (ex.: relato) → responde usuário.",
            "ai-orchestrator dispara evento para n8n (notify-n8n) com correlation ids.",
            "n8n processa/enriquece e devolve callback (n8n-callback) para atualizar o registro e/ou criar notificações.",
            "UI mostra dados enriquecidos quando disponíveis (sem travar a conversa).",
        ],
    )
    add_code_block(
        doc,
        [
            "Usuário -> ai-orchestrator (SSE)",
            "   -> (DB) cria urban_reports/transport_reports",
            "   -> notify-n8n (evento *_created)",
            "n8n -> (processa pesado / integrações)",
            "   -> n8n-callback (atualiza registro + logs)",
        ],
    )

    add_heading(doc, "9. Contratos mínimos (para não dar dor de cabeça)")
    add_bullets(
        doc,
        [
            "Idempotência: todo evento para n8n deve ter event_id (UUID) + entity_id + event_type.",
            "Retries: ai-orchestrator (curto, UX); n8n (longo, com dead-letter/alerta).",
            "Correlação: sempre propagar conversation_id, user_id, entity_id, event_id.",
            "Segurança: secrets/assinatura (x-n8n-secret) e validação no callback.",
        ],
    )

    add_heading(doc, "10. Evidências/arquivos no repositório (referência)")
    add_bullets(
        doc,
        [
            "ai-orchestrator: supabase/functions/ai-orchestrator/index.ts",
            "Integração n8n: supabase/functions/notify-n8n, n8n-webhook, n8n-callback",
            "Guia n8n: docs/N8N_INTEGRATION_GUIDE.md",
        ],
    )

    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()

