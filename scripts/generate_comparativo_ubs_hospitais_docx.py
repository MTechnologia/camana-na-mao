"""Gera docs/comparativo-ubs-hospitais-google-places.docx a partir dos números acordados no .md."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "comparativo-ubs-hospitais-google-places.docx"


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_heading(
        "Comparativo: UBS e hospitais — public_services (B × C) e Google Places",
        level=0,
    )
    t.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph(
        "Valores alinhados ao comparativo geral de equipamentos; comparação restrita "
        "aos tipos ubs e hospital."
    )

    p = doc.add_paragraph()
    p.add_run("Câmbio ilustrativo: ").bold = True
    p.add_run("R$ 5,50 por US$ 1,00.")

    p2 = doc.add_paragraph()
    p2.add_run("Google Places (estimativa): ").bold = True
    p2.add_run(
        "1 chamada Place Details (Enterprise + Atmosphere) por POI; "
        "US$ 25 / 1.000 após 1.000 grátis/mês; sem transit_station; sem Popular times."
    )

    doc.add_heading("1. Por service_type — UBS e hospital (Marcos B × C)", level=1)
    _add_table(
        doc,
        ["service_type", "B", "C", "Δ"],
        [
            ["ubs", "966", "966", "0"],
            ["hospital", "699", "699", "0"],
            ["Total (UBS + hospital)", "1.665", "1.665", "0"],
        ],
    )
    doc.add_paragraph(
        "A limpeza B → C (transporte + junk) não alterou as contagens de UBS nem de hospitais."
    )

    doc.add_heading("2. Google Places — só UBS + hospitais", level=1)
    doc.add_paragraph("POIs = soma ubs + hospital (iguais em B e C).")
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.add_run("Custo: ").bold = True
    p3.add_run("max(0, N − 1.000) × 0,025 USD; BRL = USD × 5,50.")

    _add_table(
        doc,
        ["Referência", "POIs (N)", "USD (≈)", "BRL (≈)"],
        [["UBS + hospitais (B e C)", "1.665", "16,63", "91,44"]],
    )
    doc.add_paragraph("(1.665 − 1.000) × 0,025 = 16,625 USD.")

    obs = doc.add_paragraph()
    obs.add_run("OBS: ").bold = True
    obs.add_run(
        "Para o universo completo de POIs e economia B − C no Google Places, "
        "ver o documento comparativo geral de equipamentos."
    )
    foot = doc.add_paragraph(
        "Valores ilustrativos; não substituem fatura Google."
    )
    foot.italic = True
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            continue
        if para.alignment is None:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Escrito: {OUT}")


if __name__ == "__main__":
    main()
