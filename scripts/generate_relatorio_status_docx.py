from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_kv_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"

    for k, v in items:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v


def add_roadmap_table(doc: Document) -> None:
    headers = ["Item do roadmap", "Status", "Evidências no repo", "Próximos passos / observações"]
    rows = [
        [
            "Arquitetura em React Native",
            "Parcial (PoV via WebView)",
            "docs/adr/0003-mobile-expo-react-native-com-webview-pov.md; mobile/",
            "Definir quais fluxos serão nativos (auth, navegação, formulários críticos) e planejar migração gradual.",
        ],
        [
            "Mapeamento de Integrações",
            "Parcial (documentado + functions criadas)",
            "docs/automacao_INTEGRATION_GUIDE.md; supabase/functions/*; docs/adr/0002-backend-baas-supabase.md; docs/ENVIRONMENT_VARIABLES.md",
            "Completar um diagrama único (sistemas externos + dados + credenciais) e definir responsáveis por cada integração.",
        ],
        [
            "Execução da estrutura do app em React Native",
            "Concluído (PoV)",
            "mobile/App.tsx; mobile/src/screens/FrontendWebView.tsx; docs/adr/0003-mobile-expo-react-native-com-webview-pov.md",
            "Adicionar profile 'preview' no EAS e padronizar variáveis EXPO_PUBLIC_* por ambiente (dev/hml/prod).",
        ],
        [
            "Criação da Base dos Usuários (front-end)",
            "Concluído",
            "src/pages/Register.tsx; src/contexts/AuthContext.tsx; src/pages/Profile.tsx; src/pages/profile/*",
            "Garantir consistência entre dados de auth e profiles no Supabase (migração/seed) e revisar validações de cadastro.",
        ],
        [
            "Sistema de autenticação e cadastro de usuários (front-end)",
            "Concluído",
            "src/contexts/AuthContext.tsx; src/pages/Login.tsx; src/pages/Register.tsx; src/pages/ResetPassword.tsx; src/pages/UpdatePassword.tsx",
            "Revisar redirects de confirmação de e-mail e fluxo de reset em mobile (deep links) quando migrar para nativo.",
        ],
        [
            "Perfil de usuários com dados pessoais (front-end)",
            "Concluído",
            "src/pages/profile/PersonalInfoPage.tsx; DemographicsPage.tsx; AddressPage.tsx; PreferencesPage.tsx; InterestsPage.tsx",
            "Revisar campos obrigatórios vs opcionais e alinhar com LGPD (consentimento e retenção).",
        ],
        [
            "Onboarding (front-end)",
            "Concluído",
            "src/pages/Welcome.tsx; src/pages/Onboarding.tsx; src/contexts/OnboardingContext.tsx; src/contexts/AuthContext.tsx",
            "Unificar critério de 'primeiro acesso' (profiles.onboarding_completed_at vs user_interests) para evitar divergências.",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_dir = repo_root / "docs" / "entregaveis"
    out_dir.mkdir(parents=True, exist_ok=True)

    today = date.today().isoformat()
    out_path = out_dir / f"Relatorio_Camara_na_Mao_{today}.docx"

    doc = Document()

    add_title(doc, "Câmara na Mão — Relatório Técnico (RBAC, Entregas e Roadmap)")
    add_subtitle(doc, f"Data: {today}")
    doc.add_paragraph()

    add_heading(doc, "1. Metadados do relatório")
    add_kv_table(
        doc,
        [
            ("Repositório", "camana-na-mao"),
            ("Branch (referência)", "pov-dev-mobile"),
            ("Escopo", "RBAC (Cidadão Engajado) + UX + Deploy + Migração Supabase + Roadmap"),
        ],
    )

    add_heading(doc, "2. Resumo executivo")
    doc.add_paragraph(
        "Este relatório consolida as principais entregas e correções realizadas no app, "
        "com foco em RBAC (especialmente o perfil Cidadão Engajado), experiência do usuário, "
        "integrações e deploy. Também apresenta o status dos itens de roadmap solicitados, "
        "com evidências diretamente no repositório."
    )

    add_heading(doc, "3. Entregas e melhorias implementadas (alto nível)")
    add_bullets(
        doc,
        [
            "RBAC funcional e validado no frontend (flags de permissão via useUserRole) com debug page (/debug/rbac).",
            "Menu e UI do Cidadão Engajado expõem funcionalidades habilitadas: Painéis Analíticos e Relatos (hub).",
            "Encaminhamento de relatos para vereador disponível quando permitido (Cidadão Engajado/Gestor/Admin).",
            "Audiências: listagem e inscrição integradas (fluxo interno com gravação em public.audiencia_inscricoes).",
            "Chat IA: correção de histórico misturando mensagens entre conversas (isolamento por conversationId).",
            "Agenda da Câmara: filtros por período/tipo e opção de exportar PDF.",
            "Padronização de variáveis públicas do frontend: CAMARA_URL / CAMARA_PUBLISHABLE_KEY / CAMARA_PROJECT_ID (com fallback VITE_* durante migração).",
            "Deploy do frontend no Google Cloud Run via Cloud Build (Dockerfile + Nginx, com push antes do deploy).",
            "Base de notificações com realtime (notifications + NotificationsContext) e logs de auditoria (audit_logs + tela Admin).",
        ],
    )

    add_heading(doc, "4. IA utilizada nas conversações")
    add_bullets(
        doc,
        [
            "Edge Function principal: supabase/functions/ai-orchestrator.",
            "Provedor: endpoint OpenAI-compatible configurado em AI_CHAT_BASE_URL (Supabase Secrets).",
            "Modelo: AI_CHAT_MODEL (ex.: Qwen/Qwen2.5-3B-Instruct ou equivalente no provedor).",
            "Credencial: AI_API_KEY (variável de ambiente no Supabase).",
        ],
    )

    add_heading(doc, "5. Mapeamento de integrações (visão prática)")
    add_bullets(
        doc,
        [
            "Supabase (BaaS): Auth, Postgres (RLS), Realtime, Edge Functions (docs/adr/0002-backend-baas-supabase.md).",
            "automacao: integração via Edge Functions notify-automacao e automacao-callback + guia completo em docs/automacao_INTEGRATION_GUIDE.md.",
            "Google Places: Edge Functions google-places-autocomplete e google-places-details.",
            "Agenda/Notícias/Vereadores: Edge Functions fetch-agenda, fetch-noticias, fetch-vereadores.",
        ],
    )

    add_heading(doc, "6. Roadmap — status atual (com evidências)")
    add_roadmap_table(doc)

    add_heading(doc, "7. Riscos e próximos passos recomendados")
    add_bullets(
        doc,
        [
            "Unificar a regra de onboarding (DB vs user_interests) para evitar estados divergentes.",
            "Criar workflow CI/CD para deploy de Edge Functions (não há .github/workflows no repo atualmente).",
            "Formalizar diagrama único de integrações e ambientes (dev/hml/prod), com owners e segredos por sistema.",
            "Definir plano de migração do PoV (WebView) para RN nativo em fluxos críticos.",
        ],
    )

    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()

