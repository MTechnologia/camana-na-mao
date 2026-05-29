#!/usr/bin/env python3
"""Gera relatório DOCX do fechamento 29/05/2026 — PRs #311 a #320, linguagem acessível."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "RELATORIO_FECHAMENTO_29_05_2026.docx"
DOWNLOADS_COPY = Path.home() / "Downloads" / "RELATORIO_FECHAMENTO_29_05_2026.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_status_item(doc: Document, status: str, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"[{status}] ")
    run.bold = True
    p.add_run(text)


def add_link_para(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(url)


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("M-TECH\n")
    t.bold = True
    t.font.size = Pt(22)
    t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    s = title.add_run("Relatório de fechamento diário")
    s.font.size = Pt(14)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Data de referência: {date(2026, 5, 29).strftime('%d/%m/%Y')}\n").font.size = Pt(11)
    meta.add_run("Projeto: Câmara na Mão — Assistente conversacional\n").font.size = Pt(11)
    meta.add_run("Escopo: entregas PR #311 a #320 (já em dev e homologação)\n").font.size = Pt(11)
    meta.add_run("Formato: resumo executivo e plano de testes no aplicativo").font.size = Pt(11)
    doc.add_page_break()


def add_executive_summary(doc: Document) -> None:
    add_heading(doc, "1. Resumo do dia", 1)
    add_para(
        doc,
        "Dia dedicado a concluir o pacote de melhorias do chatbot do Câmara na Mão, em dez entregas "
        "sequenciais (pull requests #311 a #320). O foco foi tornar a conversa com a assistente mais "
        "confiável, mais clara para o cidadão e mais fácil de retomar depois de fechar o aplicativo.",
    )
    add_heading(doc, "1.1 Em uma frase", 2)
    add_bullet(doc, "O assistente lembra melhor onde o cidadão parou no relato ou na avaliação.")
    add_bullet(doc, "Mensagens e botões ficaram mais estáveis (menos perda de texto, menos códigos estranhos na tela).")
    add_bullet(doc, "Relato rápido, formulários manuais e avaliação de serviços ganharam atalhos claros.")
    add_bullet(doc, "Histórico de conversas e de visitas a equipamentos pode ser limpo pelo próprio usuário.")
    add_bullet(doc, "Homologação (hml) foi alinhada com desenvolvimento (dev) ao final do dia.")


def add_camara_section(doc: Document) -> None:
    add_heading(doc, "2. Câmara na Mão — o que mudou na prática", 1)
    add_para(
        doc,
        "As entregas abaixo correspondem às pull requests #311 a #320, já integradas em dev e "
        "homologação. A descrição evita termos técnicos e foca no que o cidadão e a equipe de "
        "produto percebem ao usar o aplicativo.",
    )

    blocks = [
        (
            "Confiabilidade e acessibilidade (PR #311)",
            [
                "Respostas do assistente passam a ser anunciadas para leitores de tela quando uma nova "
                "mensagem chega.",
                "Envio de mensagem vazio ou só com espaços deixa de disparar ação indevida.",
                "Em caso de lentidão da rede, o cidadão vê aviso de tempo esgotado ou nova tentativa automática.",
                "Quando o sistema está sobrecarregado (muitas mensagens seguidas), há retentativa e depois "
                "aviso claro de limite.",
            ],
        ),
        (
            "Retomar conversa onde parou (PR #312)",
            [
                "Ao atualizar a página ou voltar depois, o assistente recupera o passo do relato ou da "
                "avaliação em andamento (tipo de relato, campos já preenchidos, barra de progresso).",
                "Perguntas mais curtas e diretas nas etapas de relato urbano e de transporte.",
                "Na etapa «quem mais é afetado» no relato urbano, aparecem botões claros: só eu, toda a rua, bairro todo.",
                "Ao trocar de assunto (por exemplo, de relato urbano para transporte), o app avisa que o "
                "progresso anterior pode ser perdido e pede confirmação.",
                "Na etapa de fotos do relato urbano, os botões de câmera e galeria aparecem no momento certo.",
            ],
        ),
        (
            "Segunda leva de refinamentos do chat (PR #313)",
            [
                "Manual completo de testes para a equipe validar o assistente passo a passo.",
                "Testes automatizados de troca de jornada (urbano ↔ transporte ↔ avaliação).",
                "Respostas mais úteis quando o modelo de IA não devolve texto (mensagem contextual em vez de erro seco).",
                "Mais tempo de espera em jornadas longas (relato e avaliação) para reduzir cortes no meio do fluxo.",
            ],
        ),
        (
            "Rótulos claros e formulário manual (PR #314)",
            [
                "Barra de progresso do relato usa os mesmos nomes que o cidadão vê nos botões (ex.: afetação "
                "«Só eu», «Toda a rua», «Bairro todo»).",
                "Atalho de formulário manual abre o formulário certo: urbano com foto ou transporte passo a passo, "
                "conforme o que estava fazendo no chat.",
                "Na avaliação de serviço, some a duplicidade de «estrelas gerais» quando já existe avaliação "
                "por dimensões (atendimento, limpeza, etc.).",
            ],
        ),
        (
            "Organização interna do aplicativo (PR #315)",
            [
                "Código do chat foi reorganizado em partes menores para facilitar manutenção futura.",
                "Para o cidadão, o comportamento permanece o mesmo; a mudança é preparação para evoluções "
                "sem instabilidade.",
            ],
        ),
        (
            "Mensagens e fotos mais seguras (PR #316)",
            [
                "Mensagens enviadas em sequência rápida deixam de sumir da conversa.",
                "Até três fotos no relato podem ser enviadas em paralelo, reduzindo tempo de espera.",
                "Botões de resposta rápida ficaram maiores e mais fáceis de usar no celular; Enter e Espaço "
                "também selecionam opção quando o foco está no botão.",
            ],
        ),
        (
            "Transporte e documentação (PR #317)",
            [
                "Parada ou ponto de ônibus deixa de ser obrigatório na barra de progresso quando o cidadão "
                "pode pular ou não souber.",
                "Link destacado no rodapé do chat para «preferir formulário manual» em relato urbano e transporte.",
                "Documentação interna sobre ordem de prioridade das respostas automáticas do assistente.",
                "Novos testes automatizados de fumaça para transporte e avaliação.",
            ],
        ),
        (
            "Relato rápido e navegação manual (PR #318)",
            [
                "Novo chip «Relato rápido» na tela inicial: fluxo urbano com menos perguntas de risco e escopo.",
                "Ao abrir formulário manual pelo chat, botão «Voltar ao chat» restaura a mesma conversa.",
                "No transporte, após informar linha, data, hora e descrição, o assistente não repete "
                "perguntas já respondidas por padrão inteligente (sentido, recorrência, impacto).",
            ],
        ),
        (
            "Correções após testes (PR #319)",
            [
                "Texto interno «[PHOTO_ATTACH_STEP]» deixa de aparecer na pergunta sobre anexar fotos.",
                "Conversa salva informações extras necessárias para retomar o fluxo no banco de dados.",
                "Ambientes que ainda não tinham a atualização do banco continuam funcionando sem travar o chat.",
            ],
        ),
        (
            "Ajustes finais validados em homologação (PR #320)",
            [
                "Durante relato urbano, se o cidadão fala de ônibus atrasado, o app oferece trocar para "
                "relato de transporte em vez de perguntar tema de buraco ou iluminação.",
                "Chip «Relato rápido» não mostra mais código técnico na bolha da mensagem.",
                "Formulário manual de transporte e urbano: seta voltar do topo retorna ao chat com a mesma conversa.",
                "Na avaliação pelo chat, rodapé «Prefiro o formulário de avaliação livre» abre a tela dedicada "
                "de avaliação com modo livre e visitas pendentes.",
                "Mensagem de sucesso do relato de transporte aparece formatada (título, campos em destaque, "
                "próximos passos), sem pular perguntas de frequência e impacto indevidamente.",
                "Tela de conversas: marcar uma ou várias, excluir selecionadas ou limpar todas; botão «+» "
                "inicia conversa nova com tela inicial limpa.",
                "Histórico de visitas a equipamentos: mesma lógica de seleção e exclusão em massa.",
            ],
        ),
    ]

    for title, items in blocks:
        add_heading(doc, title, 2)
        for item in items:
            add_bullet(doc, item)

    add_heading(doc, "2.1 O que o cidadão ganha no conjunto", 2)
    gains = [
        "Menos frustração ao relatar problema na cidade ou no transporte.",
        "Caminho alternativo claro quando prefere formulário em vez de conversa.",
        "Avaliação de UBS, escola e outros serviços mais guiada e sem telas duplicadas.",
        "Controle sobre histórico de conversas e visitas detectadas pelo aplicativo.",
        "Retomada de relato após fechar o app ou atualizar a página.",
    ]
    for g in gains:
        add_bullet(doc, g)

    add_heading(doc, "2.2 Links das entregas (GitHub)", 2)
    prs = [
        ("PR #311", "Confiabilidade e acessibilidade do chat", "https://github.com/MTechnologia/camana-na-mao/pull/311"),
        ("PR #312", "Retomar jornada (snapshot) e troca de assunto", "https://github.com/MTechnologia/camana-na-mao/pull/312"),
        ("PR #313", "Segunda onda backlog CHB + manual de testes", "https://github.com/MTechnologia/camana-na-mao/pull/313"),
        ("PR #314", "Rótulos, formulário manual e avaliação por dimensões", "https://github.com/MTechnologia/camana-na-mao/pull/314"),
        ("PR #315", "Reorganização interna do código do chat", "https://github.com/MTechnologia/camana-na-mao/pull/315"),
        ("PR #316", "Persistência de mensagens e envio de fotos", "https://github.com/MTechnologia/camana-na-mao/pull/316"),
        ("PR #317", "Transporte opcional, link manual e testes", "https://github.com/MTechnologia/camana-na-mao/pull/317"),
        ("PR #318", "Relato rápido e voltar ao chat", "https://github.com/MTechnologia/camana-na-mao/pull/318"),
        ("PR #319", "Correção de textos e banco de conversas", "https://github.com/MTechnologia/camana-na-mao/pull/319"),
        ("PR #320", "UX pós-testes (jornadas, histórico, avaliação)", "https://github.com/MTechnologia/camana-na-mao/pull/320"),
    ]
    for label, desc, url in prs:
        add_para(doc, f"{label} — {desc}")
        add_link_para(doc, "Link", url)


def add_report_types_catalog(doc: Document) -> None:
    """Catálogo de tipos de relatos e registros do cidadão (fonte: código e banco)."""
    doc.add_page_break()
    add_heading(doc, "3. Mapa de tipos de relatos e registros do cidadão", 1)
    add_para(
        doc,
        "No Câmara na Mão, o cidadão pode registrar três grandes famílias de informação "
        "estruturada (com protocolo e acompanhamento no app). Além delas, existem fluxos "
        "conversacionais que não geram «relato» no sentido administrativo (busca de serviços, "
        "audiências, dúvidas gerais). Este mapa reflete o que está implementado em maio/2026.",
    )

    add_heading(doc, "3.1 Três jornadas principais (persistidas)", 2)
    add_bullet(
        doc,
        "Relato urbano — problemas na cidade (buraco, lixo, iluminação, esgoto, etc.). "
        "Salvo na base como relato urbano; pode ser feito pelo chat, pelo chip «Relato rápido» "
        "ou pelo formulário manual com foto.",
    )
    add_bullet(
        doc,
        "Relato de transporte — ônibus, metrô e mobilidade (atraso, lotação, segurança, etc.). "
        "Salvo como relato de transporte; chat ou formulário manual passo a passo.",
    )
    add_bullet(
        doc,
        "Avaliação de serviço público — nota e comentário sobre UBS, escola, CEU e outros "
        "equipamentos. Não é «denúncia de problema na rua», mas registro formal de experiência "
        "(visitas detectadas pelo GPS ou modo livre em /avaliar).",
    )

    add_heading(doc, "3.2 Relato urbano — categorias de problema", 2)
    add_para(doc, "Ao classificar o relato, o sistema usa uma destas categorias:")
    urban_cats = [
        ("Iluminação", "Postes apagados, escuro na via."),
        ("Calçada", "Obstáculos, calçada quebrada ou inacessível."),
        ("Via pública", "Problemas gerais na via."),
        ("Pavimentação", "Buracos, asfalto, pista."),
        ("Sinalização", "Placas, semáforos, faixas."),
        ("Drenagem", "Alagamento, bueiro, enchente."),
        ("Lixo e limpeza", "Acúmulo de lixo, varrição."),
        ("Esgoto", "Vazamento, mau cheiro, saneamento."),
        ("Área verde", "Parques, poda, mato."),
        ("Higiene urbana", "Limpeza urbana ampla."),
        ("Animais", "Animais abandonados, focos, pragas."),
        ("Poluição / barulho", "Ruído, poluição ambiental."),
        ("Feedback da Câmara", "Elogio, reclamação ou sugestão dirigida à Câmara ou a vereador(a) — "
        "ver subtipos na seção 3.4."),
        ("Outro", "Quando não encaixa nas anteriores."),
    ]
    for label, hint in urban_cats:
        add_bullet(doc, f"{label}: {hint}")

    add_heading(doc, "3.2.1 Natureza do relato urbano", 2)
    add_para(doc, "Antes ou durante a coleta, o cidadão pode indicar a intenção:")
    for n in ["Reclamação", "Dúvida", "Sugestão", "Elogio"]:
        add_bullet(doc, n)

    add_heading(doc, "3.2.2 Relato urbano completo x relato rápido", 2)
    add_bullet(
        doc,
        "Relato urbano completo — todas as etapas (incluindo gravidade e «quem é afetado» "
        "quando a categoria exige).",
    )
    add_bullet(
        doc,
        "Relato rápido — fluxo enxuto pelo chip «Relato rápido»; menos perguntas de risco "
        "e escopo, para relatos simples.",
    )

    add_heading(doc, "3.3 Relato de transporte — tipos e detalhes", 2)
    add_para(doc, "O tipo principal do problema de transporte:")
    transport_types = [
        ("Atraso", "Ônibus não passou, atraso maior ou menor que 30 min, intervalo irregular."),
        ("Lotação", "Superlotação, não conseguiu embarcar, fila no ponto, ar-condicionado."),
        ("Segurança", "Assédio, furto/roubo, agressão, briga."),
        ("Acessibilidade", "Elevador/escada, rampa, veículo sem acessibilidade, assistência."),
        ("Limpeza", "Veículo sujo, mau cheiro, lixo, pragas."),
        ("Condução", "Freada brusca, imprudência, não parou no ponto."),
        ("Outro", "Demais situações — o cidadão descreve."),
    ]
    for label, hint in transport_types:
        add_bullet(doc, f"{label}: {hint}")

    add_heading(doc, "3.4 Feedback da Câmara (dentro do relato urbano)", 2)
    add_para(
        doc,
        "Quando a categoria é «Feedback da Câmara», o relato continua sendo urbano no sistema, "
        "mas o conteúdo é sobre a instituição ou parlamentares:",
    )
    for s in [
        "Elogio à Câmara ou vereador(a)",
        "Reclamação ou denúncia",
        "Sugestão",
        "Opcional: nome do vereador(a) e partido (lista de vereadores da Câmara Municipal de SP)",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "3.5 Avaliação de serviço — tipos de equipamento", 2)
    add_para(
        doc,
        "Na avaliação, o cidadão escolhe o tipo de equipamento (exemplos implementados):",
    )
    service_types = [
        "UBS", "Escola", "CEU", "Hospital", "Biblioteca", "Centro esportivo",
        "Feira", "Centro comunitário", "Creche", "CAPS", "CRAS", "Parque",
        "Assistência social", "Delegacia", "Estação de transporte", "Mercado",
        "Mercado municipal", "Teatro / cinema", "Museu", "Cemitério",
        "Acessibilidade", "Ponto de reciclagem", "Bombeiros", "Outro",
    ]
    for i in range(0, len(service_types), 4):
        chunk = ", ".join(service_types[i : i + 4])
        add_bullet(doc, chunk)

    add_heading(doc, "3.5.1 Dimensões da avaliação (quando aplicável)", 2)
    add_para(
        doc,
        "Além da nota geral (1 a 5 estrelas), o fluxo pode coletar dimensões específicas:",
    )
    for d in ["Atendimento", "Limpeza", "Infraestrutura", "Tempo de espera"]:
        add_bullet(doc, d)

    add_heading(doc, "3.6 O que não é «relato» mas o chat trata", 2)
    add_para(
        doc,
        "Estes assuntos usam o assistente, mas não geram relato urbano/transporte nem avaliação "
        "na mesma estrutura:",
    )
    for item in [
        "Serviços próximos — localizar UBS, escola, etc. no mapa",
        "Audiências públicas — consulta e inscrição",
        "Histórico do cidadão — «meus relatos», status das denúncias",
        "Informações gerais — dúvidas sobre a Câmara, legislação (base de conhecimento)",
        "Vereadores e notícias — consultas informativas",
        "Ocupação / Olho Vivo — fluxos específicos quando acionados",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.7 Resumo visual", 2)
    add_para(
        doc,
        "Relatos formais: Urbano (14 categorias + natureza + variante rápida) | Transporte "
        "(7 tipos + subdetalhes) | Avaliação (dezenas de tipos de equipamento + dimensões). "
        "Feedback da Câmara é subtipo do urbano. Formulários manuais espelham urbano e transporte.",
        bold=True,
    )


def add_report_pipeline_architecture(doc: Document) -> None:
    """Como cada jornada processa mensagens, chama LLM e usa RAG (ai-orchestrator)."""
    doc.add_page_break()
    add_heading(doc, "4. Arquitetura de processamento e uso de IA/RAG", 1)
    add_para(
        doc,
        "Cada mensagem do cidadão no chat passa pelo servidor ai-orchestrator (Supabase Edge Function). "
        "Muitas respostas são montadas por regras fixas, sem chamar o modelo de linguagem (LLM). "
        "A LLM (ex.: Gemini via Vertex) entra quando é preciso interpretar texto livre, classificar "
        "problema ou redigir resposta com contexto. O aplicativo só consome a resposta em streaming (SSE) "
        "e interpreta marcadores como [FIELD_REQUEST:...] e [COLLECTION_PROGRESS:...].",
    )

    add_heading(doc, "4.1 Pipeline comum (toda mensagem)", 2)
    steps = [
        ("1. Preparação", "Autenticação, carrega prompt de sistema (versão ativa no banco), histórico da conversa."),
        ("2. Intenção", "detectCollectionIntent — regras + palavras-chave (sem LLM): define jornada (urbano, transporte, avaliação, geral, etc.)."),
        ("3. Campos acumulados", "Lê mensagens anteriores e preenche o que já foi dito (endereço, categoria, linha…)."),
        ("4. Atalhos antes da IA", "Saudação, fora de escopo, vereador, avaliação por canal, fechamento de dúvida urbana, rotas de transporte — resposta direta, sem LLM."),
        ("5. Orquestração da coleta", "getNextMissingField — calcula o próximo campo obrigatório e o texto da pergunta (na maioria dos turnos, sem LLM)."),
        ("6. Criação automática", "Se todos os campos estão preenchidos, chama create_urban_report / create_transport_report / create_service_rating no servidor (sem LLM)."),
        ("7. Pipeline de IA", "Só se ainda precisar: monta contexto, injeta RAG se aplicável, chama LLM com ferramentas (tools), executa tool e devolve SSE."),
    ]
    for title, desc in steps:
        add_bullet(doc, f"{title}: {desc}")

    add_heading(doc, "4.2 Relato urbano — como a resposta é gerada", 2)
    for item in [
        "Turnos de coleta: o backend envia pergunta pronta com [COLLECTION_PROGRESS:urban_report:...] e [FIELD_REQUEST:campo]. O app mostra tracker, pickers (GPS, escopo, fotos).",
        "Classificação de tema: com descrição suficiente, a LLM pode chamar a ferramenta classify_report_category (categoria + subtipo em português).",
        "Endereço: validate_cep ou parsing de rua/bairro; regras evitam pedir CEP de novo se já há localização.",
        "Natureza dúvida: após descrição, busca prévia na base Supabase (searchKnowledgeBaseForUrbanDuvida) e injeta trechos no prompt; a LLM redige a resposta (search_knowledge_base é removida das tools nesse turno).",
        "Elogio / sugestão / dúvida sem reclamação: modo especial — LLM responde com instrução de não repetir coleta de risco/endereço como se fosse buraco.",
        "Feedback da Câmara (categoria feedback_camara): mesmo fluxo urbano; extração de elogio/reclamação/sugestão e vereador por regras (extractChamberFields).",
        "Relato rápido: marcador [URBAN_QUICK_REPORT] reduz perguntas de risco/escopo no orchestrator.",
        "Fim: create_urban_report grava no banco e devolve [REPORT_CREATED:id] + mensagem de sucesso (automático ou via tool da LLM).",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.3 Relato de transporte — como a resposta é gerada", 2)
    for item in [
        "Coleta guiada por getNextMissingField: tipo (report_type), subcategoria, linha, data, hora, sentido, parada, recorrência, impacto.",
        "Perguntas com pickers: [LINE_PICKER], [DATE_PICKER], [TIME_PICKER], [DIRECTION_PICKER], [SUBCATEGORY_PICKER], etc.",
        "Classificação: ferramenta classify_transport_type quando o cidadão descreve o problema com detalhe.",
        "Interceptores no servidor: antes de criar o relato, valida subcategoria inválida, parada, checklist de acessibilidade — pode devolver nova pergunta sem LLM.",
        "Modo inteligente pós-linha: após linha + data + hora + descrição, pode inferir sentido/recorrência/impacto e pular perguntas repetidas (regras + histórico).",
        "Fim: create_transport_report (automático ou tool) → mensagem formatada de sucesso no chat.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.4 Avaliação de serviço — como a resposta é gerada", 2)
    for item in [
        "Coleta de tipo de equipamento, endereço confirmado, nota (estrelas ou dimensões), comentário.",
        "Pickers: [SERVICE_TYPE_PICKER], [SERVICE_PICKER], [RATING_PICKER], [MULTI_DIMENSION_RATING_PICKER], [WAIT_TIME_PICKER].",
        "Quando completo: create_service_rating no servidor (preferência automática); falha de validação pode voltar com [FIELD_REQUEST] sem nova chamada à LLM.",
        "Atalho: mensagens vindas de canal de avaliação podem pular parte da coleta (channel rating shortcut).",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.5 Jornadas sem «relato» — LLM e ferramentas", 2)
    journeys = [
        ("Geral / tirar dúvida (general)", "Intenção general. Ver seção 4.6 sobre RAG. LLM com tools; principal: search_knowledge_base quando RAG não foi pré-injetado."),
        ("Serviços próximos (services)", "Fluxo determinístico: pergunta tipo de equipamento → CEP → find_nearby_services (mapa/GeoSampa), muitas vezes sem LLM."),
        ("Audiências (audiencias)", "Perguntas «próximas audiências»: atalho executa search_audiencias e devolve lista + [APP_ACTIONS:audiencias] sem LLM."),
        ("Histórico (history)", "Tool get_citizen_history — relatos, avaliações, inscrições."),
        ("Notícias (noticias)", "Injeta últimas notícias no prompt; LLM resume."),
        ("Vereadores (vereadores)", "Boost para general/RAG ou suggest_council_member conforme frase."),
        ("Ônibus informativo", "Hint no prompt para search_bus_lines, previsão, pontos (Olho Vivo) — não é relato de problema."),
    ]
    for title, desc in journeys:
        add_bullet(doc, f"{title}: {desc}")

    add_heading(doc, "4.6 Quando o RAG e a base de conhecimento entram", 2)
    add_para(
        doc,
        "Existem três mecanismos de «busca em base» distintos. O sistema evita duplicar: se o contexto "
        "já foi injetado no prompt, a ferramenta search_knowledge_base é removida da lista para a LLM.",
    )

    rag_rows = [
        (
            "Vertex RAG (Google)",
            "Jornada general, dúvida sobre Câmara/cidade, com VERTEX_RAG habilitado. "
            "Antes da LLM: chamada generateContent só com retrieval (datastore/corpus). "
            "Trecho injetado em [Contexto … Vertex RAG]. Não usa search_knowledge_base no mesmo turno.",
        ),
        (
            "Base Supabase (knowledge_chunks)",
            "Perguntas sobre estrutura/funcionamento interno da Câmara (general). "
            "searchKnowledgeBase pré-busca e injeta [Contexto … Supabase]. "
            "Dúvida urbana (urban_report + natureza duvida): searchKnowledgeBaseForUrbanDuvida — "
            "resposta focada no tema da cidade, sem «dump» institucional genérico.",
        ),
        (
            "Tool search_knowledge_base",
            "LLM chama durante o turno quando nenhum contexto foi pré-injetado. "
            "Usado sobretudo para zoneamento/LPUOS/SISZON, comissões, legislação urbana, "
            "e complementos quando a pré-busca não achou trecho.",
        ),
        (
            "Sem RAG",
            "Coleta estruturada (urbano/transporte/avaliação): não consulta RAG por turno — "
            "só regras, próximo campo e tools de classificação/criação. "
            "Zoneamento: pula Vertex RAG de propósito e orienta search_knowledge_base. "
            "Funcionamento da Câmara: pula Vertex, usa Supabase KB.",
        ),
    ]
    for name, desc in rag_rows:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    add_heading(doc, "4.7 Ferramentas (tools) da LLM por tipo de registro", 2)
    add_para(doc, "Ferramentas que o modelo pode invocar quando o pipeline de IA está ativo:")
    tools = [
        "classify_report_category / classify_transport_type — classificar tema (sessão, não grava ainda)",
        "validate_cep — resolver endereço por CEP",
        "create_urban_report / create_transport_report / create_service_rating — gravar registro definitivo",
        "detect_user_intent / confirm_journey_switch — troca entre jornadas estruturadas",
        "search_knowledge_base — dúvidas (se não houve injeção prévia de KB/RAG)",
        "find_nearby_services, search_audiencias, get_citizen_history — jornadas leves",
        "Ferramentas de ônibus (linhas, paradas, previsão) — consulta Olho Vivo / GeoSampa",
    ]
    for t in tools:
        add_bullet(doc, t)

    add_heading(doc, "4.8 Resumo: LLM sim ou não?", 2)
    summary_table = [
        ("Próxima pergunta do relato (campo X)", "Não — getNextMissingField + prompt fixo"),
        ("Pergunta de saudação / fora de escopo", "Não — atalho determinístico"),
        ("Listar próximas audiências", "Não — search_audiencias direto"),
        ("Serviços próximos (CEP + tipo)", "Não — fluxo services determinístico"),
        ("Classificar buraco vs iluminação", "Sim — tool classify_report_category"),
        ("Responder «o que é a Câmara?»", "Sim — com RAG Supabase ou Vertex pré-injetado"),
        ("Dúvida urbana no meio do relato", "Sim — KB Supabase injetada + redação LLM"),
        ("Criar relato com todos os campos", "Não — create_* automático no servidor"),
        ("Redigir pergunta criativa fora do script", "Sim — LLM com contexto de coleta injetado"),
    ]
    for caso, mecanismo in summary_table:
        add_bullet(doc, f"{caso} → {mecanismo}")

    add_para(
        doc,
        "Referência técnica no repositório: supabase/functions/ai-orchestrator/ (index.ts, "
        "lib-index-shortcut-order.ts, lib-index-collection-orchestration.ts, lib-index-prompt-context.ts, "
        "lib-index-ai-pipeline.ts), docs/AI_ORCHESTRATOR_SPECIFICATION.md.",
        bold=False,
    )


def add_test_plan_section(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "5. Plano de testes no aplicativo", 1)
    add_para(
        doc,
        "Roteiro para validação em ambiente de desenvolvimento ou homologação, com usuário de teste "
        "logado. Marque cada item após executar. Detalhamento técnico complementar: "
        "docs/MANUAL_TESTE_CHATBOT_COMPLETO.md no repositório.",
    )

    add_heading(doc, "5.1 Antes de começar", 2)
    add_bullet(doc, "Usar o aplicativo em dev ou hml (não produção, salvo combinação com a equipe).")
    add_bullet(doc, "Estar logado com conta de teste.")
    add_bullet(doc, "Confirmar com a equipe técnica que o assistente (servidor de IA) e o banco "
                 "de dados do ambiente estão atualizados.")

    sections = [
        (
            "Chat geral",
            [
                "Abrir a tela inicial do assistente e perguntar: «Olá, o que é a Câmara Municipal?»",
                "Esperado: resposta em linguagem clara em até cerca de 30 segundos; indicador de «digitando» some.",
                "Tentar enviar mensagem vazia: nada deve ser enviado.",
            ],
        ),
        (
            "Relato urbano (problema na cidade)",
            [
                "Tocar no chip «Relato Urbano» ou escrever que quer falar sobre a cidade.",
                "Esperado: barra «Passo X de Y» e perguntas curtas.",
                "Descrever um problema (ex.: buraco na rua) e seguir até categoria e endereço.",
                "Na pergunta de quem é afetado, usar os botões: Somente eu / Toda a rua / Bairro todo.",
                "Na etapa de fotos: aparecem Câmera e Galeria; anexar uma foto e concluir se possível.",
                "No meio do fluxo, atualizar a página (F5): progresso e campos preenchidos devem voltar coerentes.",
            ],
        ),
        (
            "Relato rápido",
            [
                "Na tela inicial, tocar em «Relato rápido».",
                "Esperado: mensagem visível só como «Quero um relato rápido» (sem códigos estranhos).",
                "Fluxo com menos perguntas de risco e escopo do que o relato urbano completo.",
            ],
        ),
        (
            "Relato de transporte",
            [
                "Iniciar jornada Transporte (chip ou texto sobre ônibus/metrô).",
                "Quando perguntar parada ou ponto, responder «pular» ou «não sei».",
                "Esperado: fluxo continua sem repetir a mesma pergunta indefinidamente.",
                "Concluir relato: mensagem de sucesso legível (título, itens em destaque, próximos passos).",
                "Confirmar que perguntas de frequência e impacto não foram puladas indevidamente.",
            ],
        ),
        (
            "Troca de jornada",
            [
                "Começar relato urbano; na descrição, escrever algo de transporte (ex.: «o ônibus está atrasando»).",
                "Esperado: aviso para trocar para relato de transporte, não pergunta de tema de buraco/iluminação.",
                "Confirmar troca: nova jornada inicia.",
                "Recusar troca em outro teste: permanece no relato urbano.",
            ],
        ),
        (
            "Formulário manual e voltar ao chat",
            [
                "No hub, chip «Formulário manual (com foto)»: abre formulário urbano sem erro na tela.",
                "Durante transporte no chat, link «Preferir formulário manual»: abre formulário passo a passo de transporte.",
                "Com conversa ativa, abrir manual e usar «Voltar ao chat» ou seta voltar: mesma conversa no assistente.",
            ],
        ),
        (
            "Avaliação de serviço",
            [
                "Chip «Avaliar serviço» na home: conversa de avaliação inicia.",
                "No rodapé, «Prefiro o formulário de avaliação livre»: abre tela /avaliar.",
                "Em /avaliar: lista de visitas pendentes (se houver) e seção de modo livre.",
                "Abrir visita válida pelo link ou botão Avaliar: serviço já contextualizado.",
                "Link de visita expirada ou já avaliada: mensagem clara e retorno à lista (sem tela branca).",
                "Modo livre: informar tipo e nome do equipamento; avaliar por dimensões sem estrelas duplicadas.",
            ],
        ),
        (
            "Histórico de conversas",
            [
                "Menu ou atalho para lista de conversas (/conversas).",
                "Abrir conversa antiga com relato em andamento: mensagens e progresso coerentes.",
                "Botão «+» ou «Nova conversa»: tela inicial limpa, sem mensagens da conversa anterior.",
                "Marcar uma ou mais conversas, excluir selecionadas; testar «Limpar todas» com confirmação.",
            ],
        ),
        (
            "Histórico de visitas",
            [
                "Perfil → Histórico de visitas.",
                "Selecionar visitas com checkbox; excluir selecionadas.",
                "Testar «Limpar todas» com confirmação (apenas em ambiente de teste).",
            ],
        ),
        (
            "Robustez",
            [
                "Com rede lenta (modo desenvolvedor do navegador): enviar mensagem; esperar aviso ou nova tentativa.",
                "Deslogar e tentar enviar: pedido de login, sem carregamento infinito.",
            ],
        ),
    ]

    for title, steps in sections:
        add_heading(doc, title, 2)
        for step in steps:
            if step.startswith("Esperado:"):
                add_bullet(doc, step)
            else:
                add_numbered(doc, step)

    add_heading(doc, "5.2 Checklist rápido pós-deploy", 2)
    smoke = [
        "Urbano: descrição → categoria → endereço → escopo → fotos",
        "Transporte: descrição → linha → conclusão ou pular parada",
        "Avaliação: visita válida + modo livre",
        "Troca de jornada com aviso",
        "Retomada após F5 com barra de progresso",
        "Nova conversa com hub limpo",
        "Exclusão em massa em conversas e visitas (ambiente de teste)",
    ]
    for s in smoke:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ ").bold = True
        p.add_run(s)


def add_status_block(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "6. Status por item (fechamento operacional)", 1)
    add_para(doc, "Visão rápida do que está pronto e do que ainda depende de validação ou deploy.")

    items = [
        ("OK", "PRs #311 a #320 integradas em desenvolvimento (dev)."),
        ("OK", "Homologação (hml) alinhada com dev no fim do dia."),
        ("OK", "Testes manuais principais executados pela equipe (jornadas, histórico, avaliação, formulários)."),
        ("OK", "Manual de testes completo disponível no repositório."),
        ("Pendente", "Deploy do front em produção (quando aprovado pelo comitê de release)."),
        ("Pendente", "Atualizações de banco em produção (metadados de conversa, exclusão de visitas pelo cidadão)."),
        ("Pendente", "Republicar assistente de IA (ai-orchestrator) em todos os ambientes após cada release."),
        ("Pendente", "Rodada formal de testes de acessibilidade com leitor de tela (opcional, recomendado)."),
    ]
    for status, text in items:
        add_status_item(doc, status, text)


def add_footer(doc: Document) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento gerado automaticamente — M-TECH — Fechamento 29/05/2026")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_cover(doc)
    add_executive_summary(doc)
    add_camara_section(doc)
    add_report_types_catalog(doc)
    add_report_pipeline_architecture(doc)
    add_test_plan_section(doc)
    add_status_block(doc)
    add_footer(doc)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    try:
        doc.save(str(OUTPUT))
        print(f"Gerado: {OUTPUT}")
    except OSError as e:
        alt = OUTPUT.with_name("RELATORIO_FECHAMENTO_29_05_2026_atualizado.docx")
        doc.save(str(alt))
        print(f"Arquivo original em uso ({e}); salvo em: {alt}")
    doc.save(str(DOWNLOADS_COPY))
    print(f"Cópia:  {DOWNLOADS_COPY}")


if __name__ == "__main__":
    main()
