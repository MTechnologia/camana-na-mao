"""
Gera dois .docx a partir de PRE_REQUISITOS_AMBIENTES.docx:
  1) Pré-requisitos até a 5.ª página (5.º salto manual de página), com coluna
     «Justificativa» nas tabelas de pré-requisitos (quadro HML x Prod e variáveis de build).
  2) Restante (a partir da 6.ª página): conteúdo sobre configuração detalhada dos ambientes.

Uso:
  py -3 scripts/build_pre_req_ambientes_docx.py

Por defeito lê e escreve em Downloads; opcionalmente copia para docs/ do repositório.
"""
from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm

# Blocos 1..58 (inclusive) = até ao 5.º w:br type=page; a partir do bloco 59 = página 6.
DOC1_BLOCK_COUNT = 58

# Tabelas na primeira parte: índices 0=componentes, 1=HMLxProd, 2=variáveis build
TABLE_IDX_ADD_JUSTIFICATION = (1, 2)

JUSTIFICATIONS_HML_PROD = [
    "Permite validar merges e pipelines numa branch dedicada sem afetar a linha acessível ao público.",
    "Homologação pode usar plano e capacidade menores; produção exige projeto dimensionado para disponibilidade, backup e suporte contratual.",
    "Cold start em homologação reduz custo fixo; em produção, instâncias mínimas evitam lentidão em picos de utilização.",
    "Teto baixo em HML contém custos e carga acidental; em produção o teto acompanha picos com autoscaling.",
    "Recursos reduzidos em HML bastam para testes; produção precisa de margem para tráfego real, gzip e utilizadores simultâneos.",
    "Subdomínio separado evita conflito de cookies, cache e DNS entre validação interna e o site público.",
    "A mesma solução gerenciada simplifica operação e renovação nos dois ambientes.",
    "Retenção curta em HML limita armazenamento; produção exige janela maior e PITR para recuperação granular e conformidade.",
    "Retenção alinhada ao risco: menor em HML; maior em produção para auditoria e incidentes (LGPD).",
    "Homologação sem SLA formal; produção com meta explícita alinhada ao contrato de serviço.",
    "HML pode ser atualizada a qualquer momento; produção concentra mudanças em horário de menor impacto para o cidadão.",
    "HML partilha quota para experimentação; produção com quota e alertas dedicados reduz surpresas de fatura.",
    "Prazo de recuperação mais longo é aceitável em HML; produção exige restauração mais rápida após incidentes.",
    "Em HML admite-se maior perda de dados recentes; em produção o PITR reduz a perda máxima aceitável.",
]

JUSTIFICATIONS_BUILD_VARS = [
    "URL base da API Supabase injetada no bundle; o SPA deve apontar para o projeto correto em cada ambiente.",
    "Chave anon do Supabase usada no cliente com RLS; identifica o tenant sem expor chaves privilegiadas.",
    "Referência curta do projeto (ref) usada em integrações e diagnóstico do pipeline de deploy.",
    "Habilita mapas e geocodificação na interface; deve ser restrita por domínio/referrer para reduzir abuso.",
    "Par público VAPID para Web Push no browser; sem esta chave as notificações web ficam desativadas ou degradadas.",
]


def _strip_page_breaks_from_paragraph(p_el) -> None:
    for run in p_el.findall(qn("w:r")):
        for br in list(run.findall(qn("w:br"))):
            if br.get(qn("w:type")) == "page":
                run.remove(br)


def _remove_trailing_page_breaks(doc: DocumentObject) -> None:
    body = doc.element.body
    if not len(body):
        return
    last = body[-1]
    if isinstance(last, CT_P):
        _strip_page_breaks_from_paragraph(last)


def _append_body_slice(
    target: DocumentObject, elements: list, start: int, end: int | None
) -> None:
    slice_ = elements[start:end]
    for el in slice_:
        target.element.body.append(deepcopy(el))


def _new_doc_from_body_slice(
    elements: list, start: int, end: int | None
) -> DocumentObject:
    out = Document()
    body = out.element.body
    # Documento vazio do python-docx só traz w:sectPr; o conteúdo deve ficar antes dele.
    sect_pr = body[-1]
    body.remove(sect_pr)
    _append_body_slice(out, elements, start, end)
    body.append(sect_pr)
    return out


def _add_justification_column(table, header: str, texts: list[str]) -> None:
    table.add_column(Cm(4.5))
    rows = table.rows
    rows[0].cells[-1].text = header
    for i, row in enumerate(rows[1:], start=0):
        if i < len(texts):
            row.cells[-1].text = texts[i]
        else:
            row.cells[-1].text = ""


def build(
    source: Path,
    out_prereq: Path,
    out_config: Path,
) -> None:
    src = Document(source)
    body_children = list(src.element.body)

    doc1 = _new_doc_from_body_slice(body_children, 0, DOC1_BLOCK_COUNT)
    _remove_trailing_page_breaks(doc1)

    for ti in TABLE_IDX_ADD_JUSTIFICATION:
        t = doc1.tables[ti]
        if ti == 1:
            _add_justification_column(
                t, "Justificativa", JUSTIFICATIONS_HML_PROD
            )
        elif ti == 2:
            _add_justification_column(
                t, "Justificativa", JUSTIFICATIONS_BUILD_VARS
            )

    doc1.save(out_prereq)

    doc2 = _new_doc_from_body_slice(body_children, DOC1_BLOCK_COUNT, None)
    doc2.save(out_config)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        type=Path,
        default=Path(r"C:\Users\felip\Downloads\PRE_REQUISITOS_AMBIENTES.docx"),
    )
    parser.add_argument(
        "--out-prereq",
        type=Path,
        default=Path(r"C:\Users\felip\Downloads\PRE_REQUISITOS_AMBIENTES.docx"),
    )
    parser.add_argument(
        "--out-config",
        type=Path,
        default=Path(
            r"C:\Users\felip\Downloads\PRE_REQUISITOS_AMBIENTES_CONFIGURACAO_ATUAL.docx"
        ),
    )
    parser.add_argument(
        "--also-docs",
        action="store_true",
        help="Copiar também para docs/ no repositório (caminho fixo do projeto).",
    )
    args = parser.parse_args()

    if not args.source.is_file():
        raise SystemExit(f"Ficheiro de origem inexistente: {args.source}")

    backup = args.source.with_suffix(".antes_split.docx")
    if args.out_prereq.resolve() == args.source.resolve():
        import shutil

        shutil.copy2(args.source, backup)
        print(f"Backup: {backup}")

    build(args.source, args.out_prereq, args.out_config)
    print(f"Pré-requisitos (até ~pág. 5): {args.out_prereq}")
    print(f"Configuração atual (pág. 6+): {args.out_config}")

    if args.also_docs:
        import shutil

        repo_docs = Path(__file__).resolve().parents[1] / "docs"
        repo_docs.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.out_prereq, repo_docs / args.out_prereq.name)
        shutil.copy2(args.out_config, repo_docs / args.out_config.name)
        print(f"Cópias em: {repo_docs}")


if __name__ == "__main__":
    main()
