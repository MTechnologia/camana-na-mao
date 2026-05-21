"""Gera relatório .docx das PRs #270 e #271 (tasks NREF)."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "RELATORIO_NREF_PR270_PR271.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("Relatório de entregas — tasks NREF", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Câmara na Mão — Perfil, visitas e autenticação")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)

    meta = doc.add_paragraph(f"Data do relatório: {date.today().strftime('%d/%m/%Y')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_heading(doc, "Visão geral", 1)
    doc.add_paragraph(
        "Este documento consolida as tasks de referência (NREF) entregues nas Pull Requests "
        "#270 e #271, com orientações de validação em tela (smoke test) para QA e homologação."
    )

    add_heading(doc, "Pull Requests", 2)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "PR"
    hdr[1].text = "Branch"
    hdr[2].text = "Status / base"
    table.rows[1].cells[0].text = "#270"
    table.rows[1].cells[1].text = "feature/nref-perfil-fixes"
    table.rows[1].cells[2].text = "Mergeada em dev"
    table.rows[2].cells[0].text = "#271"
    table.rows[2].cells[1].text = "feature/nref012-020-fixes"
    table.rows[2].cells[2].text = "Aberta → dev"
    doc.add_paragraph()
    doc.add_paragraph("URLs:")
    add_bullet(doc, " https://github.com/MTechnologia/camana-na-mao/pull/270")
    add_bullet(doc, " https://github.com/MTechnologia/camana-na-mao/pull/271")

    add_heading(doc, "Pré-requisitos gerais (PR #271)", 1)
    prereq = [
        "Aplicar migrations: 20260524120000_expire_visits_set_departed_at.sql e "
        "20260524120100_backfill_service_visits_departed_at.sql",
        "Deploy da Edge Function send-email (--no-verify-jwt) e secret APP_PUBLIC_URL "
        "(URL pública do app, ex.: Cloud Run ou app-mtechnologia.com)",
        "Supabase → Authentication → URL Configuration: incluir https://<domínio>/nova-senha "
        "nas Redirect URLs",
        "Front-end na branch feature/nref012-020-fixes ou dev após merge do PR #271",
    ]
    for item in prereq:
        add_bullet(doc, item)

    # --- PR 270 ---
    add_heading(doc, "PR #270 — NREF006, NREF007, NREF011", 1)
    doc.add_paragraph(
        "Foco: Minhas Inscrições, integração de interesses com alertas de audiências e "
        "esclarecimento de favoritos vs. lembretes."
    )

    tasks_270 = [
        (
            "NREF006 — Título Minhas Inscrições",
            "Alinhamento do título da página ao item do menu do perfil.",
            [
                "Acessar /perfil/inscricoes",
                "Verificar título da página: Minhas Inscrições (não variações antigas)",
            ],
        ),
        (
            "NREF007 — Interesses e alertas de audiências",
            "Interesses do perfil passam a alimentar alertas/recomendações de audiências; "
            "política RLS em audiencia_topic_alerts; central de inscrições simplificada.",
            [
                "Perfil → Interesses: salvar temas de interesse",
                "Verificar alertas/recomendações de audiências conforme temas",
                "Minhas Inscrições: abas Serviços, Transporte, Alertas e temas, Audiências",
            ],
        ),
        (
            "NREF011 — Meus Favoritos (coração vs. sino)",
            "Copy comparativa entre favoritar equipamento e inscrever-se em lembretes/notificações.",
            [
                "Perfil → card Meus Favoritos ou /servicos/favoritos",
                "Conferir texto explicando diferença entre coração (favorito) e sino (alertas/lembretes)",
            ],
        ),
    ]

    for title, desc, steps in tasks_270:
        add_heading(doc, title, 2)
        doc.add_paragraph(desc)
        doc.add_paragraph("Como testar em tela:", style="Intense Quote")
        for i, step in enumerate(steps, 1):
            add_numbered(doc, step)
        doc.add_paragraph()

    # --- PR 271 ---
    add_heading(doc, "PR #271 — NREF012 a NREF020", 1)
    doc.add_paragraph(
        "Foco: histórico de visitas (saída), padronização de UI do perfil, lembretes de audiências, "
        "limpeza de preferências, remoção de menus redundantes, e-mail de recovery e acessibilidade em estudo."
    )

    tasks_271 = [
        (
            "NREF012 — Saída no histórico de visitas",
            "departed_at preenchido ao avaliar, dispensar ou expirar visita (≥48h), não apenas ao sair do GPS.",
            [
                "Ativar detecção em Perfil → Preferências → Localização e visitas",
                "Gerar visita pendente (geofence / Perto de você)",
                "Avaliar ou dispensar → Perfil → Histórico de visitas: conferir horário de saída",
                "Visita expirada (>48h): conferir departed_at após RPC/job de expiração",
            ],
        ),
        (
            "NREF013 — Padronização de textos do perfil",
            "Títulos alinhados ao menu (ex.: Preferências em vez de Configurações).",
            [
                "Navegar /perfil e subpáginas (dados pessoais, preferências, inscrições, direitos, visitas)",
                "Conferir título no header de cada tela",
            ],
        ),
        (
            "NREF014 — Calendários e relógios padronizados",
            "Componentes standard-date-picker, standard-time-input e standard-calendar-panel.",
            [
                "Perfil → Dados demográficos: seletor de data de nascimento",
                "Perto de você / filtros com data: mesmo padrão visual",
                "Relato de transporte: seletor de horário padronizado",
            ],
        ),
        (
            "NREF015 — CTA lembretes (Minhas Inscrições → Audiências)",
            "Empty state com botão para explorar audiências; com inscrições, Inscrever em mais lembretes.",
            [
                "/perfil/inscricoes?aba=audiencias sem lembretes: texto + CTA para /audiencias",
                "Inscrever em audiência na ficha do evento → voltar à aba: lista + botão extra",
            ],
        ),
        (
            "NREF016 — Acessibilidade (estudo em andamento)",
            "Card Acessibilidade removido do perfil; rota /configuracoes/acessibilidade informativa.",
            [
                "/perfil: não deve listar Acessibilidade",
                "/configuracoes/acessibilidade: página de estudo + Rever tutorial",
                "Admin /admin/settings/accessibility: inalterado",
            ],
        ),
        (
            "NREF017 — Campos removidos em Preferências",
            "Removidos Newsletter e seção Privacidade; detecção de visitas em card Localização e visitas.",
            [
                "/perfil/preferencias: sem Newsletter e sem Privacidade",
                "Card Localização e visitas com toggle de detecção",
                "Salvar preferências: toast de sucesso",
            ],
        ),
        (
            "NREF018 — Botão de lembretes em Preferências",
            "Preferências explica canais; link Ver minhas inscrições em audiências.",
            [
                "Card Lembretes de audiências: sem CTA de inscrever para /audiencias",
                "Botão abre /perfil/inscricoes?aba=audiencias",
            ],
        ),
        (
            "NREF019 — Exportar Dados fora do menu do perfil",
            "Card Exportar Dados removido de Meu Perfil; exportação permanece em Direitos LGPD.",
            [
                "/perfil: sem card Exportar Dados",
                "/perfil/direitos: exportação disponível",
                "/perfil/exportar-dados: rota direta ainda funciona",
            ],
        ),
        (
            "NREF020 — Link único no e-mail de redefinição de senha",
            "E-mail com link direto para /nova-senha?token_hash=...; app usa verifyOtp.",
            [
                "Login → Esqueci minha senha → receber e-mail (após deploy send-email)",
                "Botão e link de fallback com a mesma URL do app",
                "Abrir link → Criar nova senha → login com nova senha",
                "Link inválido/expirado: mensagem e redirecionamento para novo pedido",
            ],
        ),
    ]

    for title, desc, steps in tasks_271:
        add_heading(doc, title, 2)
        doc.add_paragraph(desc)
        doc.add_paragraph("Como testar em tela:", style="Intense Quote")
        for step in steps:
            add_numbered(doc, step)
        doc.add_paragraph()

    add_heading(doc, "Roteiro sugerido (smoke test ~30 min)", 1)
    smoke = [
        "Perfil: menus sem Exportar Dados e sem Acessibilidade",
        "Preferências: layout novo + link para inscrições em audiências",
        "Minhas Inscrições → Audiências: empty state e CTAs",
        "Dados demográficos + filtro com data: calendário padronizado",
        "Esqueci minha senha: e-mail + nova senha (com send-email deployado)",
        "Visita → avaliar ou dispensar → histórico com saída registrada",
    ]
    for step in smoke:
        add_numbered(doc, step)

    add_heading(doc, "Fora do escopo do PR #271", 1)
    doc.add_paragraph(
        "Alterações locais não incluídas no commit: CompleteInvitePage.tsx, "
        "delete-user e migration 20260521180000_prepare_auth_user_deletion.sql."
    )

    add_heading(doc, "Deploy pós-merge (PR #271)", 1)
    deploy = [
        "npx supabase db push (ou aplicar migrations no ambiente alvo)",
        "npx supabase functions deploy send-email --no-verify-jwt",
        "Configurar APP_PUBLIC_URL nos secrets da Edge Function",
    ]
    for item in deploy:
        add_bullet(doc, item)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Gerado: {OUTPUT}")


if __name__ == "__main__":
    main()
